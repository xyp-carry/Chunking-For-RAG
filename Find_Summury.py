from docx import Document
import pandas as pd
import numpy as np
import re
from tqdm import tqdm

# Find all the tables of document
def extract_tables_from_docx(filepath):
    doc = Document(filepath)
    tables_data = []

    for table in doc.tables:
        table_data = []
        for row in table.rows:
            row_data = [cell.text for cell in row.cells]
            table_data.append(row_data)
        tables_data.append(table_data)

    return tables_data
def extract_text_from_docx(docx_path):
    # 打开.docx文件
    doc = Document(docx_path)
    text = []
    
    # 遍历文档中的每个段落
    for paragraph in doc.paragraphs:
        text.append(paragraph.text)
    
    # 返回所有段落的文本，合并为一个字符串
    return '\n'.join(text)

# Determine whether it is a page number.
def if_array_asc(array):
    last = -1
    for i in array.index:
        if last == -1:
            last = int(array[i])
            continue
        if int(array[i]) >= last:
            last = int(array[i])
        else:
            return False
    return True

# Build meta data about document structure
def build_layer(array_index, array_content, array_page):
    max_layer = 1
    layer_dict = {'content':[]}
    select_layer = layer_dict
    for idx in array_index.index:
        if array_index[idx] == '' or array_index[idx] == np.nan or array_index[idx] == '\t':
            layer_dict['content'].append((array_content[idx],array_page[idx]))
        else:
            degree = re.split('[-_.|]',array_index[idx].replace(' ',''))
            if len(degree) > max_layer:
                max_layer = len(degree)
            comb_d = ''
            for i,d in enumerate(degree):
                if comb_d == '':
                    comb_d += d
                else: 
                    comb_d += f'-{d}'
                if not select_layer.get(comb_d, False):
                    select_layer[comb_d] = {'content':[(array_content[idx],array_page[idx])]}
                    select_layer = layer_dict
                else:
                    select_layer = select_layer[comb_d]
    layer_dict['max_layer'] = max_layer
    return layer_dict


def find_directory(tables, trigger = False):
    for i, table in enumerate(tables):
        table_dict = {}
        col_key = []
        for col in table[0]:
            table_dict[col] = []
            col_key.append(col)
        for row in table[1:]:
            for idx, word in enumerate(row):
                table_dict[col_key[idx]].append(word)
        try:
            df = pd.DataFrame()
            for col in table_dict:
                df[col] = table_dict[col]
            for col in df.columns:
                if 'page' in col.lower() and if_array_asc(df[col]):
                    layer = build_layer(df[df.columns[0]],df[df.columns[1]],df[df.columns[2]])
        except Exception as e:
            break
    return layer 

# get block of document
def get_block(text:str, title_list:list):
    block = []
    use_keys_list = []
    for title in title_list:
        if isinstance(title[0],list):
            parts = text.split(title[0][0][0],1)
            if len(parts) < 2:
                continue
            use_keys_list.append(title[1])
            block.append(parts[0])
            text = parts[1]
        else:
            parts = text.split(title[0],1)
            if len(parts) < 2:
                continue
            use_keys_list.append(title[1])
            block.append(parts[0])
            text = parts[1]
    block.append(text)
    return block[1:], use_keys_list

# get every layer's key by Recursion
def text_split(text:str, layer:dict, res_document:list):
    title_list = []
    for key in layer.keys():
        if key not in ['content','max_layer']:
            title_list.append((layer[key]['content'],key))
    # print(title_list)
    block, use_keys_list = get_block(text, title_list)
    if len(block) == 0:
        res_document.append(text)
    for idx in range(len(block)):
        text_split(block[idx], layer[use_keys_list[idx]], res_document)



if __name__ == '__main__':
    filepath = 'test.docx'

    tables = extract_tables_from_docx(filepath)
    layer = find_directory(tables)

    text = extract_text_from_docx(filepath)
    res_document = []
    text_split(text, layer, res_document)
    print(res_document)
    for table in tables:
        dict_tables = {"item":[]}
        for row in table[1:]:
            dict_table = {}
            for idx in range(len(row)):
                dict_table[table[0][idx]] = row[idx]
            dict_tables['item'].append(dict_table)
        res_document.append(dict_tables)
    print(res_document)