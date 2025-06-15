from docx import Document
import pandas as pd
import numpy as np
import re
from tqdm import tqdm


def transform_num(string):
    trans_dict= {'一':1,'二':2,'三':3,'四':4,'五':5,'六':6,'七':7,'八':8,'九':9,'十':1}
    new_str = ''
    for s in string:
        if s in trans_dict:
            new_str+= str(trans_dict[s])
    if new_str !='':
        return int(new_str)
    else:
        return False

def build_layer(array_index, array_content, array_page):

    max_layer = 1
    layer_dict = {'content':[]}
    select_layer = layer_dict
    for idx in array_index.index:
        if array_index[idx] == '' or array_index[idx] == np.nan or array_index[idx] == '\t':
            layer_dict['content'].append((array_content[idx],array_page[idx]))
        else:
            degree = re.split('[-_.|]',array_index[idx].replace(' ',''))
            if degree[-1] == '':
                degree = degree[:-1]
            if len(degree) > max_layer:
                max_layer = len(degree)
            comb_d = ''
            for i,d in enumerate(degree):
                if comb_d == '':
                    comb_d += d
                else: 
                    comb_d += f'-{d}'
                if not select_layer.get(comb_d, False):
                    select_layer[comb_d] = {'content':[(array_content[idx],array_page[idx])]}
                    select_layer = layer_dict
                else:
                    select_layer = select_layer[comb_d]
    layer_dict['max_layer'] = max_layer
    return layer_dict


def get_table_to_layer(tables_data):
    for i, table in enumerate(tables_data):
        layer = False
        table_dict = {}
        col_key = []
        for col in table[0]:
            table_dict[col] = []
            col_key.append(col)
        # print(col_key)
        for row in table[1:]:
            for idx, word in enumerate(row):
                table_dict[col_key[idx]].append(word)
        try:
            df = pd.DataFrame()
            for col in table_dict:
                df[col] = table_dict[col]
            if 'page' in df.columns[2].lower() and len(list(df.columns)) == 3:
                layer = build_layer(df[df.columns[0]],df[df.columns[1]],df[df.columns[2]])
                return layer
        except Exception as e:
            print(e)
            return False
    return False

def find_directory_sdt(filepath):
    ns = {
    'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    }
    doc = Document(filepath)
    TOC_tables = []
# 查询所有内容控制元素
    for sdt in doc.element.xpath('.//w:sdt'):
        # 查询内容控制元素的标签
        TOC_table = [['Code', 'Content', 'Page']]
        sdtContent = sdt.find('.//w:sdtContent', namespaces=ns)
        if sdtContent is not None:
            for Content in sdtContent:
                text_list = []
                for text1 in Content.xpath('.//w:t'):
                    text_list.append(text1.text)
                if len(text_list) == 3:
                    TOC_table.append([text_list[0], text_list[1], text_list[2]])
            TOC_tables.append(TOC_table)

    layer = get_table_to_layer(TOC_tables)

    return layer

def find_directory_toc(filepath):
    TOC_table = [['Code', 'Content', 'Page']]
    doc = Document(filepath)
    a= 1
    for i in doc.paragraphs:
        toc = []
        if i.style.name.startswith('toc'):
            list_toc = i.text.split("\t")

            match = re.search(r'^-?[0-9.]+', list_toc[0])
            max_page = 0
            if match and match.start() == 0:
                TOC_table.append([str(match.group()), list_toc[0],list_toc[1]])
                max_page = list_toc[1]
            elif transform_num(list_toc[0].split(' ')[0]):
                TOC_table.append([str(transform_num(list_toc[0].split(' ')[0])), list_toc[0],list_toc[1]])
                max_page = list_toc[1]
            else:
                if list_toc[1].isdigit():
                    TOC_table.append(['',list_toc[0],list_toc[1]])
                else:
                    TOC_table.append(['',list_toc[0],str(max_page)])

    layer = get_table_to_layer([TOC_table])
    
    return layer

def find_directory_table(filepath):
    doc = Document(filepath)
    tables_data = []

    for table in doc.tables:
        table_data = []
        for row in table.rows:
            row_data = [cell.text for cell in row.cells]
            table_data.append(row_data)
        tables_data.append(table_data)
    
    layer = get_table_to_layer(tables_data)

    return layer 


def File_reader(file_path):
    """
    The function Get FILE and return chunk and layer about chunk.
    First, the function read FILE and find directory .
    Second, 

    """
    extension = file_path.split('.')[-1]
    if extension == 'docx':
        layer = find_directory_table(file_path)
        if layer == False:
            layer = find_directory_sdt(file_path)
        if layer == False:
            layer = find_directory_toc(file_path)
    return layer



print(File_reader(r'.\FILE\test_d.docx'))